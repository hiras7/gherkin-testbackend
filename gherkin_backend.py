
from flask import Flask, request, send_file, jsonify, make_response
from flask_cors import CORS
from docx import Document
import os, re, time
from collections import defaultdict

app = Flask(__name__)

FRONTEND_ORIGIN = os.environ.get("FRONTEND_ORIGIN", "*")
CORS(app, resources={r"/*": {"origins": FRONTEND_ORIGIN}})

UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "outputs"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

BLOCK_HEADER_RE = re.compile(r'^\[(?P<ref>[^\]]+)\]\s+(?P<title>.+)$')

def digits_from_ref(refcode: str) -> str:
    m = re.findall(r'(\d+)', refcode or '')
    return m[-1] if m else 'UNKNOWN'

def norm_heading(text: str) -> str:
    t = (text or '').strip().lower()
    if t.endswith(':'):
        t = t[:-1].strip()
    return t

def para_is_heading(text: str, label: str) -> bool:
    return norm_heading(text) == label

def is_fit_heading(text: str) -> bool:
    t = norm_heading(text)
    return t in ('fit criteria', 'fitcriterion', 'fit-criteria', 'fit', 'acceptance criteria', 'acceptance tests')

def parse_requirements_from_docx(path: str):
    doc = Document(path)
    items, cur, section = [], None, None

    def commit():
        nonlocal cur
        if cur:
            cur.setdefault('ReferenceCode', '')
            cur.setdefault('ReqID', digits_from_ref(cur.get('ReferenceCode')))
            cur.setdefault('Title', '')
            cur.setdefault('ReqName', f"[{cur.get('ReferenceCode','')}] {cur.get('Title','')}".strip())
            cur.setdefault('Requirement', '')
            cur.setdefault('Rationale', '')
            cur.setdefault('FitCriteria', [])
            items.append(cur)
            cur = None

    for p in doc.paragraphs:
        text = (p.text or '').strip()
        if not text:
            continue
        m = BLOCK_HEADER_RE.match(text)
        if m:
            commit()
            cur = {'ReferenceCode': m.group('ref').strip(), 'Title': m.group('title').strip(), 'FitCriteria': []}
            section = None
            continue
        if para_is_heading(text, 'requirement'):
            section = 'Requirement'
            if cur:
                cur.setdefault('Requirement', '')
            continue
        if para_is_heading(text, 'rationale') or para_is_heading(text, 'rational'):
            section = 'Rationale'
            if cur:
                cur.setdefault('Rationale', '')
            continue
        if is_fit_heading(text):
            section = 'Fit'
            continue
        if not cur:
            continue
        if section == 'Requirement':
            cur['Requirement'] = (cur.get('Requirement', '') + (' ' if cur.get('Requirement') else '') + text)
        elif section == 'Rationale':
            cur['Rationale'] = (cur.get('Rationale', '') + (' ' if cur.get('Rationale') else '') + text)
        elif section == 'Fit':
            cur['FitCriteria'].append(text)
    commit()
    return [{
        'ReqID': it.get('ReqID'),
        'ReqName': it.get('ReqName'),
        'Topic': it.get('Title'),
        'Requirement': it.get('Requirement'),
        'Rationale': it.get('Rationale'),
        'FitCriteria': it.get('FitCriteria', [])
    } for it in items]

THEME_SPLIT_RE = re.compile(r'\s*[:\-–—>→]\s*')

def extract_theme(line: str):
    s = (line or '').strip()
    if not s:
        return None
    parts = THEME_SPLIT_RE.split(s, maxsplit=1)
    if len(parts) > 1:
        theme = parts[0].strip().lower()
        theme = re.sub(r'^[\-\*•]+\s*', '', theme)
        return theme or None
    return None

def group_fits_by_theme(fits):
    groups = defaultdict(list)
    for line in fits or []:
        theme = extract_theme(line) or 'misc'
        groups[theme].append(line)
    return dict(groups)

def scenario_count_by_mode(fits, mode: str) -> int:
    if mode == 'atomized':
        return len(fits or [])
    if mode == 'ultra-optimized':
        return 1
    n = len(fits or [])
    themes = group_fits_by_theme(fits)
    multiple_topics = len(themes.keys()) > 1
    if n <= 3:
        return 1
    if 4 <= n <= 10 and not multiple_topics:
        return 2
    return 3

def actor_from_text(requirement_text: str, flags) -> str:
    if flags.get('opt_strict_actor'):
        m = re.search(r'As a[n]?\s+([A-Za-z0-9 _/\-]+)', requirement_text or '', re.IGNORECASE)
        if m:
            return m.group(1).strip()
    return 'the user'

def distribute_into_buckets(items, k: int):
    if k <= 1:
        return [items]
    buckets = [[] for _ in range(k)]
    for i, it in enumerate(items or []):
        buckets[i % k].append(it)
    return buckets

def themed_buckets(fits, k: int):
    groups = group_fits_by_theme(fits)
    sorted_groups = sorted(groups.items(), key=lambda x: len(x[1]), reverse=True)
    if len(sorted_groups) >= k:
        picks = [list(x) for x in sorted_groups[:k]]
        leftovers = [ln for _, ls in sorted_groups[k:] for ln in ls]
        for i, ln in enumerate(leftovers):
            picks[i % k][1].append(ln)
        return [(name, lines) for name, lines in picks]
    rr = distribute_into_buckets(fits, k)
    return [(f'group {i+1}', rr[i]) for i in range(k)]

def compute_overview(data, mode):
    out = []
    for r in data:
        fits = r.get('FitCriteria') or []
        out.append({'ReqID': r.get('ReqID'), 'ReqName': r.get('ReqName'), 'FitCount': len(fits), 'ScenarioCount': scenario_count_by_mode(fits, mode)})
    return out

def compute_overview_totals(data, mode):
    total_requirements = len(data or [])
    total_fits = sum(len(r.get('FitCriteria') or []) for r in data or [])
    total_scenarios = sum(scenario_count_by_mode(r.get('FitCriteria') or [], mode) for r in data or [])
    return {'totalRequirements': total_requirements, 'totalFitCriteria': total_fits, 'totalScenarios': total_scenarios}

def build_rules(mode, flags):
    rules = []
    if mode == 'atomized':
        rules.append('Mode: Atomized — each FIT criterion becomes its own scenario')
    elif mode == 'ultra-optimized':
        rules.append('Mode: Ultra-Optimized — one scenario per requirement (legacy behavior)')
    else:
        rules.append('Mode: Optimized — ≤3→1, 4–10→2, >10 or multi-topic→3 scenarios')
    rules.append(f"Scenario Outline Optimization: {'ON' if flags.get('opt_outline') else 'OFF'}")
    rules.append(f"Preserve Bullet Formatting: {'ON' if flags.get('opt_preserve_bullets') else 'OFF'}")
    rules.append(f"Strict Actor Referencing: {'ON' if flags.get('opt_strict_actor') else 'OFF'}")
    rules.extend(['Gherkin v46 style; third-person actors; no OR in steps; no UI implementation details', 'Given the user is logged in (unless an explicit different actor is detected)'])
    return rules

def build_traceability(data, mode: str, top_n: int = 20):
    ranked = sorted(data, key=lambda r: len(r.get('FitCriteria') or []), reverse=True)
    if top_n and top_n > 0:
        ranked = ranked[:top_n]
    nodes, links, seen = [], [], set()
    def add_node(name, type_, label=None):
        if name in seen:
            return
        nodes.append({'name': name, 'type': type_, 'label': label or name})
        seen.add(name)
    for r in ranked:
        req_id = r.get('ReqID', 'UNKNOWN')
        topic = (r.get('Topic') or r.get('ReqName') or '').strip() or req_id
        fits = r.get('FitCriteria') or []
        req_node = f'REQ:{req_id}'
        add_node(req_node, 'req', label=topic)
        k = scenario_count_by_mode(fits, mode)
        if k == 1:
            sc_name = f'SC:{req_id}:1'
            add_node(sc_name, 'sc', label=f'{req_id} — SC1')
            links.append({'source': req_node, 'target': sc_name, 'value': max(1, len(fits))})
            themes = group_fits_by_theme(fits)
            for th, items in themes.items():
                theme_node = f'TH:{th}'
                add_node(theme_node, 'theme', label=th)
                links.append({'source': sc_name, 'target': theme_node, 'value': max(1, len(items))})
        else:
            buckets = themed_buckets(fits, k)
            for idx, (name, group) in enumerate(buckets, 1):
                sc_name = f'SC:{req_id}:{idx}'
                label = name if name and name != 'misc' else f'Group {idx}'
                add_node(sc_name, 'sc', label=f'{req_id} — {label}')
                links.append({'source': req_node, 'target': sc_name, 'value': max(1, len(group))})
                themes = group_fits_by_theme(group)
                for th, items in themes.items():
                    theme_node = f'TH:{th}'
                    add_node(theme_node, 'theme', label=th)
                    links.append({'source': sc_name, 'target': theme_node, 'value': max(1, len(items))})
    return {'nodes': nodes, 'links': links, 'meta': {'topN': top_n, 'reqs': len(ranked)}}

def ts_identifier(s: str) -> str:
    s = (s or '').strip()
    s = re.sub(r'[^\w\s\-]', '', s)
    s = re.sub(r'\s+', ' ', s)
    return s[:80]

def data_testid_from_text(s: str) -> str:
    base = re.sub(r'[^A-Za-z0-9]+', '-', (s or '').strip().lower()).strip('-')
    return base or 'element'

def render_ts_block_for_requirement(r: dict, mode: str) -> str:
    req_id = r.get('ReqID', 'UNKNOWN')
    topic = r.get('Topic', '') or r.get('ReqName', '')
    fits = r.get('FitCriteria', []) or []
    n_scen = scenario_count_by_mode(fits, mode)
    lines = []
    feature_name = ts_identifier(f'{req_id} {topic}'.strip())
    lines.append(f"test.describe('{feature_name}', () => {{")
    if mode == 'atomized' and fits:
        for i, fit in enumerate(fits, 1):
            scen_name = ts_identifier(f'{topic} — FIT {i}')
            tid = data_testid_from_text(fit)
            lines.extend([
                f"  test('{scen_name}', async ({{ page }}) => {{",
                f"    await page.goto(`${{BASE_URL}}/dashboard`); // pre-auth assumed",
                f"    // FIT: {fit}",
                f"    await expect(page.getByTestId('{tid}')).toBeVisible();",
                f"  }});"
            ])
    else:
        k = n_scen
        if k == 1:
            scen_name = ts_identifier(topic or req_id)
            lines.extend([
                f"  test('{scen_name}', async ({{ page }}) => {{",
                f"    await page.goto(`${{BASE_URL}}/dashboard`); // pre-auth assumed",
            ])
            if fits:
                for j, fit in enumerate(fits):
                    tid = data_testid_from_text(fit)
                    prefix = 'Then' if j == 0 else 'And'
                    lines.append(f"    // {prefix}: {fit}")
                    lines.append(f"    await expect(page.getByTestId('{tid}')).toBeVisible();")
            else:
                lines.append('    // Then: it should meet the specified acceptance criteria')
            lines.append('  });')
        else:
            buckets = themed_buckets(fits, k)
            for idx, (name, group) in enumerate(buckets, 1):
                scen_name = ts_identifier(f"{topic} — {name or f'Group {idx}'}")
                lines.extend([
                    f"  test('{scen_name}', async ({{ page }}) => {{",
                    f"    await page.goto(`${{BASE_URL}}/dashboard`); // pre-auth assumed",
                ])
                if group:
                    for j, fit in enumerate(group):
                        tid = data_testid_from_text(fit)
                        prefix = 'Then' if j == 0 else 'And'
                        lines.append(f"    // {prefix}: {fit}")
                        lines.append(f"    await expect(page.getByTestId('{tid}')).toBeVisible();")
                else:
                    lines.append('    // Then: it should meet the specified acceptance criteria')
                lines.append('  });')
    lines.append('});')
    return '\n'.join(lines)

def generate_playwright_ts(input_path: str, output_ts_path: str, mode: str = 'optimized'):
    data = parse_requirements_from_docx(input_path)
    chunks = [
        '// Auto-generated by Gherkin Intelligence Engine',
        "import { test, expect } from '@playwright/test';",
        '',
        "const BASE_URL = process.env.BASE_URL ?? 'https://your-app.example.com';",
        ''
    ]
    for r in data:
        chunks.append(render_ts_block_for_requirement(r, mode))
        chunks.append('')
    ts_code = '\n'.join(chunks)
    with open(output_ts_path, 'w', encoding='utf-8') as f:
        f.write(ts_code)
    return True

@app.route('/', methods=['GET'])
def root():
    return jsonify({'service': 'gherkin-backend', 'status': 'ok'}), 200

@app.route('/healthz', methods=['GET'])
def healthz():
    return jsonify({'ok': True}), 200

def get_mode_flags_guidelines(form):
    mode = (form.get('mode') or 'optimized').strip().lower()
    if mode not in ('ultra-optimized', 'optimized', 'atomized'):
        mode = 'optimized'
    flags = {
        'opt_outline': form.get('opt_outline') == '1',
        'opt_preserve_bullets': form.get('opt_preserve_bullets') == '1',
        'opt_strict_actor': form.get('opt_strict_actor') == '1',
    }
    guidelines = (form.get('guidelines') or '').strip()
    return mode, flags, guidelines

@app.route('/preview', methods=['POST'])
def preview():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file.filename == '' or not file.filename.lower().endswith('.docx'):
        return jsonify({'error': 'Invalid file (.docx expected)'}), 400
    path = os.path.join(UPLOAD_FOLDER, file.filename)
    file.save(path)
    mode, flags, guidelines = get_mode_flags_guidelines(request.form)
    try:
        top_n = int((request.form.get('topN') or '20').strip())
    except Exception:
        top_n = 20
    t0 = time.perf_counter()
    data = parse_requirements_from_docx(path)
    overview = compute_overview(data, mode)
    totals = compute_overview_totals(data, mode)
    trace = build_traceability(data, mode, top_n=top_n)
    elapsed = round(time.perf_counter() - t0, 3)
    return jsonify({'time': elapsed, 'data': data, 'rules': build_rules(mode, flags), 'overview': overview, 'overviewTotals': totals, 'traceability': trace, 'guidelines': guidelines})

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file.filename == '' or not file.filename.lower().endswith('.docx'):
        return jsonify({'error': 'Invalid file (.docx expected)'}), 400
    path = os.path.join(UPLOAD_FOLDER, file.filename)
    file.save(path)
    mode, flags, guidelines = get_mode_flags_guidelines(request.form)
    out = os.path.join(OUTPUT_FOLDER, 'gherkin_output.docx')
    t0 = time.perf_counter()
    _ = generate_gherkin_document(path, out, mode=mode, flags=flags, guidelines=guidelines)
    elapsed = round(time.perf_counter() - t0, 3)
    resp = make_response(send_file(out, as_attachment=True))
    resp.headers['X-Process-Time'] = str(elapsed)
    return resp

@app.route('/generate_playwright', methods=['POST'])
def generate_playwright():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file.filename == '' or not file.filename.lower().endswith('.docx'):
        return jsonify({'error': 'Invalid file (.docx expected)'}), 400
    path = os.path.join(UPLOAD_FOLDER, file.filename)
    file.save(path)
    mode, _, _ = get_mode_flags_guidelines(request.form)
    out_ts = os.path.join(OUTPUT_FOLDER, 'gherkin_tests.spec.ts')
    t0 = time.perf_counter()
    _ = generate_playwright_ts(path, out_ts, mode=mode)
    elapsed = round(time.perf_counter() - t0, 3)
    resp = make_response(send_file(out_ts, as_attachment=True))
    resp.headers['X-Process-Time'] = str(elapsed)
    return resp

def generate_gherkin_document(input_path, output_path, mode='optimized', flags=None, guidelines=''):
    flags = flags or {}
    data = parse_requirements_from_docx(input_path)
    doc = Document()
    for r in data:
        req_id = r.get('ReqID', 'UNKNOWN')
        req_name = r.get('ReqName', '')
        topic = r.get('Topic', '') or r.get('ReqName', '')
        requirement = r.get('Requirement', '')
        rationale = r.get('Rationale', '')
        fits = r.get('FitCriteria', [])
        doc.add_paragraph(f'REQ ID: {req_id}')
        doc.add_paragraph(f'REQ NAME: {req_name}')
        doc.add_paragraph('')
        doc.add_paragraph(f'Feature: {topic}')
        actor = actor_from_text(requirement, flags)
        doc.add_paragraph(f'As a {actor}')
        doc.add_paragraph(f'I want {topic.lower()}')
        doc.add_paragraph(f'So that {rationale or "business value is achieved"}')
        doc.add_paragraph('')
        mode_local = mode
        if mode_local == 'atomized' and fits:
            for i, fit in enumerate(fits, 1):
                doc.add_paragraph(f'@REQ-{req_id}')
                doc.add_paragraph(f'Scenario: {topic} — FIT {i}')
                doc.add_paragraph('Given the user is logged in' if actor == 'the user' else f'Given {actor} is authenticated')
                doc.add_paragraph(f'When the system evaluates requirement {req_id}')
                doc.add_paragraph(f'Then {fit}')
                doc.add_paragraph('')
        else:
            k = scenario_count_by_mode(fits, mode_local)
            if k == 1:
                doc.add_paragraph(f'@REQ-{req_id}')
                doc.add_paragraph(f'Scenario: {topic}')
                doc.add_paragraph('Given the user is logged in' if actor == 'the user' else f'Given {actor} is authenticated')
                doc.add_paragraph(f'When the system evaluates requirement {req_id}')
                if fits:
                    doc.add_paragraph(f'Then it should satisfy {len(fits)} FIT criteria')
                    for line in fits:
                        doc.add_paragraph(f'And {line}')
                else:
                    doc.add_paragraph('Then it should meet the specified acceptance criteria')
                doc.add_paragraph('')
            else:
                buckets = themed_buckets(fits, k)
                for idx, (name, lines) in enumerate(buckets, 1):
                    doc.add_paragraph(f'@REQ-{req_id}')
                    suffix = f' — {name.title()}' if name and name != 'misc' else f' — Group {idx}'
                    doc.add_paragraph(f'Scenario: {topic}{suffix}')
                    doc.add_paragraph('Given the user is logged in' if actor == 'the user' else f'Given {actor} is authenticated')
                    doc.add_paragraph(f'When the system evaluates requirement {req_id}')
                    if lines:
                        for j, line in enumerate(lines):
                            if j == 0:
                                doc.add_paragraph(f'Then {line}')
                            else:
                                doc.add_paragraph(f'And {line}')
                    else:
                        doc.add_paragraph('Then it should meet the specified acceptance criteria')
                    doc.add_paragraph('')
    def append_meta(document):
        document.add_paragraph('Rules Applied', style='Heading 1')
        for rr in build_rules(mode, flags):
            document.add_paragraph(f'- {rr}')
        if guidelines:
            document.add_paragraph('Guidelines (provided)', style='Heading 1')
            for line in guidelines.splitlines():
                document.add_paragraph(line)
    doc.save(output_path)
    doc2 = Document(output_path)
    append_meta(doc2)
    doc2.add_paragraph('Summary Table', style='Heading 1')
    tbl = doc2.add_table(rows=1, cols=5, style='Table Grid')
    hdr = ['Topic', 'Req ID', 'Name', '# FIT Criteria', '# Gherkin Scenarios']
    for i, h in enumerate(hdr):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    parsed = parse_requirements_from_docx(input_path)
    for r in parsed:
        fits = r.get('FitCriteria', [])
        row = tbl.add_row().cells
        row[0].text = r.get('Topic', '') or ''
        row[1].text = r.get('ReqID', '') or ''
        row[2].text = r.get('ReqName', '') or ''
        row[3].text = str(len(fits))
        row[4].text = str(scenario_count_by_mode(fits, mode))
    doc2.save(output_path)
    return True

if __name__ == '__main__':
    port = int(os.environ.get('PORT', '5000'))
    app.run(host='0.0.0.0', port=port, debug=False)
